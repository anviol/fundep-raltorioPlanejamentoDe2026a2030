"""Biblioteca de apoio para gerar os planos de ação em .docx.

Reproduz em Word o mesmo visual dos arquivos .tex do repositório:
papel timbrado da Fundep como fundo de todas as páginas, rodapé
institucional na base, cabeçalho "Página N" a partir da segunda
página, fonte Times New Roman e a paleta de cores do documento
original do Word.

Depende de python-docx. As imagens vêm de img/.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell

RAIZ = Path(__file__).resolve().parent.parent
IMG = RAIZ / "img"
TIMBRADO = IMG / "timbrado_a4_fundep.png"
RODAPE = IMG / "rodape_fundep.jpg"

# Cores do documento Word, iguais às definidas nos arquivos .tex
COR = {
    "azulTitulo": "1F4E79",     # títulos de seção e linhas de atividade
    "azulAtividade": "2F5496",  # títulos das atividades 01..NN
    "azulClaro": "D9E2F3",      # coluna esquerda da tabela de identificação
    "azulCabecalho": "9CC2E5",  # cabeçalho das tabelas
    "cinzaClaro": "F2F2F2",     # zebra das linhas do resumo
    "azulVivo": "0070C0",       # anos, TOTAL do cronograma e links
    "azulCelula": "C0E6F5",     # linha de trimestres do cronograma
    "cinzaMedio": "BFBFBF",     # linhas de atividade do cronograma
    "cinzaE8": "E8E8E8",        # células de valores do cronograma
    "bordaCinza": "BFBFBF",     # bordas internas
    "bordaEscura": "595959",    # bordas externas
    "branco": "FFFFFF",
}

# Proporção largura/altura do timbrado (2481 x 3508 px)
RAZAO_TIMBRADO = 2481 / 3508

# Namespaces necessários para o VML usado no timbrado e no rodapé
NS_VML = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:w10="urn:schemas-microsoft-com:office:word"'
)

# Definição do tipo de forma de imagem do VML, igual à do Word
SHAPETYPE = (
    '<v:shapetype id="_x0000_t75" coordsize="21600,21600" filled="f" stroked="f"'
    ' o:spt="75" o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe">'
    '<v:stroke joinstyle="miter"/>'
    "<v:formulas>"
    '<v:f eqn="if lineDrawn pixelLineWidth 0"/><v:f eqn="sum @0 1 0"/>'
    '<v:f eqn="sum 0 0 @1"/><v:f eqn="prod @2 1 2"/>'
    '<v:f eqn="prod @3 21600 pixelWidth"/><v:f eqn="prod @3 21600 pixelHeight"/>'
    '<v:f eqn="sum @0 0 1"/><v:f eqn="prod @6 1 2"/>'
    '<v:f eqn="prod @7 21600 pixelWidth"/><v:f eqn="sum @8 21600 0"/>'
    '<v:f eqn="prod @7 21600 pixelHeight"/><v:f eqn="sum @10 21600 0"/>'
    "</v:formulas>"
    '<v:path gradientshapeok="t" o:connecttype="rect" o:extrusionok="f"/>'
    '<o:lock v:ext="edit" aspectratio="t"/>'
    "</v:shapetype>"
)

_contador_forma = [1024]


def _proximo_spid():
    _contador_forma[0] += 1
    return "_x0000_s%d" % _contador_forma[0]


# ------------------------------------------------------------------
# Documento, seções, timbrado e rodapé
# ------------------------------------------------------------------

def novo_documento(espaco_paragrafo_pt=12, entrelinha=1.5):
    """Cria o documento com a fonte, o espaçamento e o papel do plano."""
    doc = Document()

    normal = doc.styles["Normal"]
    _fonte(normal, "Times New Roman", Pt(12))
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(espaco_paragrafo_pt)
    pf.line_spacing = entrelinha
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nome in ("Header", "Footer"):
        estilo = doc.styles[nome]
        _fonte(estilo, "Times New Roman", Pt(10))
        estilo.paragraph_format.space_after = Pt(0)
        estilo.paragraph_format.line_spacing = 1.0
        estilo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    secao = doc.sections[0]
    configura_secao(secao, paisagem=False, primeira_pagina_diferente=True)
    return doc


def _fonte(estilo, nome, tamanho):
    estilo.font.name = nome
    estilo.font.size = tamanho
    rpr = estilo.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), nome)


def nova_secao_paisagem(doc):
    """Abre uma seção em paisagem, como a quebra de seção do Word."""
    from docx.enum.section import WD_SECTION

    secao = doc.add_section(WD_SECTION.NEW_PAGE)
    configura_secao(secao, paisagem=True, primeira_pagina_diferente=False)
    return secao


def configura_secao(secao, paisagem, primeira_pagina_diferente):
    """Define papel, margens, cabeçalho de página e o timbrado da seção."""
    if paisagem:
        secao.orientation = WD_ORIENT.LANDSCAPE
        secao.page_width = Cm(29.7)
        secao.page_height = Cm(21)
        secao.top_margin = Cm(2)
        secao.bottom_margin = Cm(3)
    else:
        secao.orientation = WD_ORIENT.PORTRAIT
        secao.page_width = Cm(21)
        secao.page_height = Cm(29.7)
        secao.top_margin = Cm(2)
        secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(2)
    secao.right_margin = Cm(2)
    secao.header_distance = Cm(1.25)
    secao.footer_distance = Cm(1.0)
    secao.different_first_page_header_footer = primeira_pagina_diferente

    cabecalhos = [secao.header]
    rodapes = [secao.footer]
    if primeira_pagina_diferente:
        cabecalhos.append(secao.first_page_header)
        rodapes.append(secao.first_page_footer)

    for indice, cabecalho in enumerate(cabecalhos):
        cabecalho.is_linked_to_previous = False
        paragrafo = cabecalho.paragraphs[0]
        paragrafo.text = ""
        formata_paragrafo(
            paragrafo, alinhamento="esquerda", espaco_depois=0, entrelinha=1.0
        )
        # A primeira página não numera, como no documento original
        if indice == 0:
            escreve(paragrafo, "Página ", tamanho=10)
            _campo_pagina(paragrafo)
        _timbrado(paragrafo, paisagem)

    for rodape in rodapes:
        rodape.is_linked_to_previous = False
        paragrafo = rodape.paragraphs[0]
        paragrafo.text = ""
        formata_paragrafo(
            paragrafo, alinhamento="esquerda", espaco_depois=0, entrelinha=1.0
        )
        _rodape_institucional(paragrafo, paisagem)


def _campo_pagina(paragrafo):
    """Insere o campo PAGE, que o Word atualiza sozinho."""
    inicio = paragrafo.add_run()
    marca = OxmlElement("w:fldChar")
    marca.set(qn("w:fldCharType"), "begin")
    inicio._r.append(marca)

    instrucao_run = paragrafo.add_run()
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = "PAGE   \\* MERGEFORMAT"
    instrucao_run._r.append(instrucao)

    separador_run = paragrafo.add_run()
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    separador_run._r.append(separador)

    paragrafo.add_run("2")

    fim_run = paragrafo.add_run()
    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")
    fim_run._r.append(fim)


def _timbrado(paragrafo, paisagem):
    """Papel timbrado como imagem de fundo, atrás do texto.

    Em retrato o timbrado cobre a página inteira. Em paisagem ele é
    encaixado pela altura e alinhado à direita, para que a faixa com
    "fundep.ufmg.br" fique na borda direita sem distorcer a imagem,
    que é um A4 em retrato.
    """
    id_relacao, _ = paragrafo.part.get_or_add_image(str(TIMBRADO))
    if paisagem:
        altura = 595.3
        largura = altura * RAZAO_TIMBRADO
        horizontal = "right"
    else:
        largura, altura = 595.3, 841.9
        horizontal = "center"

    xml = (
        '<w:pict %s>%s'
        '<v:shape id="timbrado%s" o:spid="%s" type="#_x0000_t75"'
        ' style="position:absolute;margin-left:0;margin-top:0;'
        "width:%.2fpt;height:%.2fpt;z-index:-251658240;"
        "mso-position-horizontal:%s;mso-position-horizontal-relative:page;"
        'mso-position-vertical:center;mso-position-vertical-relative:page"'
        ' o:allowincell="f">'
        '<v:imagedata r:id="%s" o:title="timbrado_a4"/>'
        '<w10:wrap anchorx="page" anchory="page"/>'
        "</v:shape></w:pict>"
    ) % (
        NS_VML,
        SHAPETYPE,
        _contador_forma[0],
        _proximo_spid(),
        largura,
        altura,
        horizontal,
        id_relacao,
    )
    run = paragrafo.add_run()
    run._r.append(parse_xml(xml))


def _rodape_institucional(paragrafo, paisagem):
    """Cobre a faixa de rodapé do timbrado e aplica o rodapé atual.

    O timbrado traz um rodapé antigo, que é coberto por um retângulo
    branco de 2,3 cm na base, como nos arquivos .tex, e substituído
    pela imagem de rodapé centralizada.
    """
    largura_pagina = 841.9 if paisagem else 595.3
    largura_rodape = 510.0  # 18 cm em pontos
    altura_rodape = largura_rodape * _proporcao_rodape()

    id_relacao, _ = paragrafo.part.get_or_add_image(str(RODAPE))

    tarja = (
        '<w:pict %s>'
        '<v:rect id="tarja%s" o:spid="%s"'
        ' style="position:absolute;margin-left:0;margin-top:0;'
        "width:%.2fpt;height:65.2pt;z-index:-251658230;"
        "mso-position-horizontal:center;mso-position-horizontal-relative:page;"
        'mso-position-vertical:bottom;mso-position-vertical-relative:page"'
        ' fillcolor="#ffffff" stroked="f" o:allowincell="f">'
        '<w10:wrap anchorx="page" anchory="page"/>'
        "</v:rect></w:pict>"
    ) % (NS_VML, _contador_forma[0], _proximo_spid(), largura_pagina)

    imagem = (
        '<w:pict %s>%s'
        '<v:shape id="rodape%s" o:spid="%s" type="#_x0000_t75"'
        ' style="position:absolute;margin-left:0;margin-top:0;'
        "width:%.2fpt;height:%.2fpt;z-index:-251658220;"
        "mso-position-horizontal:center;mso-position-horizontal-relative:page;"
        'mso-position-vertical:bottom;mso-position-vertical-relative:page"'
        ' o:allowincell="f">'
        '<v:imagedata r:id="%s" o:title="rodape_fundep"/>'
        '<w10:wrap anchorx="page" anchory="page"/>'
        "</v:shape></w:pict>"
    ) % (
        NS_VML,
        SHAPETYPE,
        _contador_forma[0],
        _proximo_spid(),
        largura_rodape,
        altura_rodape,
        id_relacao,
    )

    for xml in (tarja, imagem):
        run = paragrafo.add_run()
        run._r.append(parse_xml(xml))


_proporcao = []


def _proporcao_rodape():
    """Altura dividida pela largura da imagem de rodapé."""
    if not _proporcao:
        from docx.image.image import Image

        imagem = Image.from_file(str(RODAPE))
        _proporcao.append(imagem.px_height / imagem.px_width)
    return _proporcao[0]


# ------------------------------------------------------------------
# Parágrafos
# ------------------------------------------------------------------

def paragrafo(
    doc,
    texto="",
    tamanho=12,
    negrito=False,
    italico=False,
    cor=None,
    alinhamento="justificado",
    espaco_antes=None,
    espaco_depois=None,
    entrelinha=None,
    recuo_esquerda=None,
    recuo_deslocado=None,
):
    """Insere um parágrafo com a formatação indicada."""
    par = doc.add_paragraph()
    formata_paragrafo(
        par,
        alinhamento=alinhamento,
        espaco_antes=espaco_antes,
        espaco_depois=espaco_depois,
        entrelinha=entrelinha,
        recuo_esquerda=recuo_esquerda,
        recuo_deslocado=recuo_deslocado,
    )
    if texto:
        escreve(par, texto, tamanho=tamanho, negrito=negrito, italico=italico, cor=cor)
    return par


ALINHAMENTOS = {
    "justificado": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "esquerda": WD_ALIGN_PARAGRAPH.LEFT,
    "centro": WD_ALIGN_PARAGRAPH.CENTER,
    "direita": WD_ALIGN_PARAGRAPH.RIGHT,
}


def formata_paragrafo(
    par,
    alinhamento=None,
    espaco_antes=None,
    espaco_depois=None,
    entrelinha=None,
    recuo_esquerda=None,
    recuo_deslocado=None,
):
    pf = par.paragraph_format
    if alinhamento:
        pf.alignment = ALINHAMENTOS[alinhamento]
    if espaco_antes is not None:
        pf.space_before = Pt(espaco_antes)
    if espaco_depois is not None:
        pf.space_after = Pt(espaco_depois)
    if entrelinha is not None:
        pf.line_spacing = entrelinha
    if recuo_esquerda is not None:
        pf.left_indent = Cm(recuo_esquerda)
    if recuo_deslocado is not None:
        pf.first_line_indent = Cm(-recuo_deslocado)
    return par


def escreve(par, texto, tamanho=12, negrito=False, italico=False, cor=None, fonte=None):
    """Acrescenta um trecho de texto ao parágrafo."""
    run = par.add_run(texto)
    run.font.size = Pt(tamanho)
    run.bold = negrito
    run.italic = italico
    if cor:
        run.font.color.rgb = RGBColor.from_string(COR.get(cor, cor))
    nome = fonte or "Times New Roman"
    run.font.name = nome
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), nome)
    return run


def titulo_documento(doc, principal, secundario):
    """Título centralizado da primeira página."""
    par = paragrafo(
        doc, alinhamento="centro", espaco_depois=6, entrelinha=1.0
    )
    escreve(par, principal, tamanho=17, negrito=True, cor="azulTitulo")
    par = paragrafo(
        doc, alinhamento="centro", espaco_depois=12, entrelinha=1.0
    )
    escreve(par, secundario, tamanho=13, negrito=True, cor="azulTitulo")


def titulo_secao(doc, texto):
    """Título de seção: 13pt, negrito, azul escuro."""
    par = paragrafo(
        doc, alinhamento="esquerda", espaco_antes=16, espaco_depois=8, entrelinha=1.0
    )
    escreve(par, texto, tamanho=13, negrito=True, cor="azulTitulo")
    # O título nunca fica sozinho no fim da página
    par.paragraph_format.keep_with_next = True
    return par


def titulo_atividade(doc, texto):
    """Título de atividade: 12pt, negrito, azul médio."""
    par = paragrafo(
        doc, alinhamento="esquerda", espaco_antes=16, espaco_depois=8, entrelinha=1.0
    )
    escreve(par, texto, tamanho=12, negrito=True, cor="azulAtividade")
    par.paragraph_format.keep_with_next = True
    return par


def item_lista(doc, marcador, texto, recuo=1.25, tamanho=12, espaco_depois=3, entrelinha=None):
    """Item de lista com marcador manual e recuo deslocado."""
    par = paragrafo(
        doc,
        alinhamento="justificado",
        espaco_antes=0,
        espaco_depois=espaco_depois,
        entrelinha=entrelinha,
        recuo_esquerda=recuo,
        recuo_deslocado=0.55,
    )
    escreve(par, "%s\t" % marcador, tamanho=tamanho)
    escreve(par, texto, tamanho=tamanho)
    _tabulacao(par, recuo)
    return par


def _tabulacao(par, posicao_cm):
    """Marca de tabulação para alinhar o texto após o marcador."""
    ppr = par._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(Cm(posicao_cm).twips)))
    tabs.append(tab)


def hiperlink(par, texto, url, tamanho=12, cor="azulVivo"):
    """Insere um hiperlink no parágrafo."""
    id_relacao = par.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), id_relacao)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    for tag, valor in (("w:color", COR[cor]), ("w:sz", str(int(tamanho * 2)))):
        elemento = OxmlElement(tag)
        elemento.set(qn("w:val"), valor)
        rpr.append(elemento)
    fontes = OxmlElement("w:rFonts")
    for atributo in ("w:ascii", "w:hAnsi", "w:cs"):
        fontes.set(qn(atributo), "Times New Roman")
    rpr.insert(0, fontes)
    run.append(rpr)
    texto_xml = OxmlElement("w:t")
    texto_xml.set(qn("xml:space"), "preserve")
    texto_xml.text = texto
    run.append(texto_xml)
    link.append(run)
    par._p.append(link)


def quebra_pagina(doc):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run()
    quebra = OxmlElement("w:br")
    quebra.set(qn("w:type"), "page")
    run._r.append(quebra)
    return par


# ------------------------------------------------------------------
# Tabelas
# ------------------------------------------------------------------

def nova_tabela(
    doc,
    larguras_cm,
    borda_cor="bordaCinza",
    borda_pt=0.5,
    borda_externa_cor=None,
    margem_celula_cm=0.19,
    margem_vertical_cm=0.06,
    alinhamento="centro",
):
    """Cria uma tabela de largura fixa, com bordas e margens definidas."""
    tabela = doc.add_table(rows=0, cols=len(larguras_cm))
    tabela.autofit = False
    tabela.alignment = {
        "centro": WD_TABLE_ALIGNMENT.CENTER,
        "esquerda": WD_TABLE_ALIGNMENT.LEFT,
    }[alinhamento]
    tabela._larguras = list(larguras_cm)

    tbl_pr = tabela._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    largura_total = OxmlElement("w:tblW")
    largura_total.set(qn("w:w"), str(int(Cm(sum(larguras_cm)).twips)))
    largura_total.set(qn("w:type"), "dxa")
    tbl_pr.append(largura_total)

    _bordas_tabela(tbl_pr, borda_cor, borda_pt, borda_externa_cor)

    margens = OxmlElement("w:tblCellMar")
    for lado, valor in (
        ("top", margem_vertical_cm),
        ("left", margem_celula_cm),
        ("bottom", margem_vertical_cm),
        ("right", margem_celula_cm),
    ):
        elemento = OxmlElement("w:%s" % lado)
        elemento.set(qn("w:w"), str(int(Cm(valor).twips)))
        elemento.set(qn("w:type"), "dxa")
        margens.append(elemento)
    tbl_pr.append(margens)

    grade = tabela._tbl.find(qn("w:tblGrid"))
    for coluna, largura in zip(grade.findall(qn("w:gridCol")), larguras_cm):
        coluna.set(qn("w:w"), str(int(Cm(largura).twips)))
    return tabela


def _bordas_tabela(tbl_pr, cor, espessura_pt, cor_externa=None):
    bordas = OxmlElement("w:tblBorders")
    tamanho = str(max(2, int(round(espessura_pt * 8))))
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        externa = lado in ("top", "left", "bottom", "right")
        elemento = OxmlElement("w:%s" % lado)
        elemento.set(qn("w:val"), "single")
        elemento.set(qn("w:sz"), tamanho)
        elemento.set(qn("w:space"), "0")
        escolhida = cor_externa if (externa and cor_externa) else cor
        elemento.set(qn("w:color"), COR.get(escolhida, escolhida))
        bordas.append(elemento)
    tbl_pr.append(bordas)


def linha(tabela, celulas, alinhamento_vertical="topo", altura_minima_cm=None):
    """Acrescenta uma linha a partir da descrição das células.

    Cada célula é um dicionário com as chaves opcionais: texto, span,
    fundo, negrito, italico, cor, tamanho, alinhamento, vertical e
    entrelinha. A soma dos spans precisa fechar o número de colunas.
    """
    nova = tabela.add_row()
    if altura_minima_cm:
        nova.height = Cm(altura_minima_cm)
    tr = nova._tr
    posicao = 0
    coluna = 0
    for descricao in celulas:
        span = descricao.get("span", 1)
        tcs = tr.findall(qn("w:tc"))
        tc = tcs[posicao]
        if span > 1:
            largura = sum(tabela._larguras[coluna:coluna + span])
            _define_largura(tc, largura)
            grid_span = OxmlElement("w:gridSpan")
            grid_span.set(qn("w:val"), str(span))
            tc.tcPr.insert(0, grid_span)
            for extra in tcs[posicao + 1:posicao + span]:
                tr.remove(extra)
        else:
            _define_largura(tc, tabela._larguras[coluna])
        _preenche(
            _Cell(tc, tabela),
            descricao,
            alinhamento_vertical=descricao.get("vertical", alinhamento_vertical),
        )
        posicao += 1
        coluna += span
    return nova


def _define_largura(tc, largura_cm):
    tc_pr = tc.get_or_add_tcPr()
    antiga = tc_pr.find(qn("w:tcW"))
    if antiga is not None:
        tc_pr.remove(antiga)
    largura = OxmlElement("w:tcW")
    largura.set(qn("w:w"), str(int(Cm(largura_cm).twips)))
    largura.set(qn("w:type"), "dxa")
    tc_pr.append(largura)


VERTICAIS = {
    "topo": WD_ALIGN_VERTICAL.TOP,
    "centro": WD_ALIGN_VERTICAL.CENTER,
    "base": WD_ALIGN_VERTICAL.BOTTOM,
}


def _preenche(celula, descricao, alinhamento_vertical="topo"):
    celula.vertical_alignment = VERTICAIS[alinhamento_vertical]
    fundo = descricao.get("fundo")
    if fundo:
        sombreado = OxmlElement("w:shd")
        sombreado.set(qn("w:val"), "clear")
        sombreado.set(qn("w:color"), "auto")
        sombreado.set(qn("w:fill"), COR.get(fundo, fundo))
        celula._tc.get_or_add_tcPr().append(sombreado)

    par = celula.paragraphs[0]
    formata_paragrafo(
        par,
        alinhamento=descricao.get("alinhamento", "esquerda"),
        espaco_antes=descricao.get("espaco_antes", 0),
        espaco_depois=descricao.get("espaco_depois", 0),
        entrelinha=descricao.get("entrelinha", 1.0),
    )
    texto = descricao.get("texto", "")
    if texto:
        escreve(
            par,
            texto,
            tamanho=descricao.get("tamanho", 10),
            negrito=descricao.get("negrito", False),
            italico=descricao.get("italico", False),
            cor=descricao.get("cor"),
        )
    return celula


def mantem_tabelas_inteiras(doc):
    """Impede que as tabelas se dividam entre páginas.

    Cada linha ganha cantSplit, e os parágrafos de todas as linhas,
    menos os da última, ganham keepNext. Com isso o Word leva a tabela
    inteira para a página seguinte quando ela não cabe no espaço
    restante, como acontece no PDF gerado pelo LaTeX.
    """
    for tabela in doc.tables:
        linhas = tabela._tbl.findall(qn("w:tr"))
        for indice, tr in enumerate(linhas):
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = OxmlElement("w:trPr")
                tr.insert(0, tr_pr)
            tr_pr.append(OxmlElement("w:cantSplit"))
            if indice == len(linhas) - 1:
                continue
            for par in tr.iter(qn("w:p")):
                ppr = par.find(qn("w:pPr"))
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    par.insert(0, ppr)
                ppr.insert(0, OxmlElement("w:keepNext"))


def celula(texto="", **atributos):
    """Atalho para descrever uma célula."""
    atributos["texto"] = texto
    return atributos


def vazias(quantidade, **atributos):
    """Sequência de células vazias, útil no cronograma."""
    return [dict(atributos) for _ in range(quantidade)]
